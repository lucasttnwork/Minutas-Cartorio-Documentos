#!/usr/bin/env python3
"""
extract_with_gemini.py - Fase 3: Extracao Contextual com Gemini 3 Flash

Este script processa documentos de cartorio usando o Gemini 3 Flash para
interpretacao contextual avancada DIRETAMENTE (sem OCR intermediario).
Usa a capacidade multimodal do Gemini para processar PDFs e imagens:

1. REESCRITA: Transcrever o documento de forma organizada
2. EXPLICACAO CONTEXTUAL: Interpretar semanticamente o documento
3. DADOS CATALOGADOS: Extrair JSON estruturado

Uso:
    python extract_with_gemini.py FC_515_124_p280509
    python extract_with_gemini.py FC_515_124_p280509 --limit 5
    python extract_with_gemini.py FC_515_124_p280509 --type MATRICULA_IMOVEL
    python extract_with_gemini.py FC_515_124_p280509 --verbose

Autor: Pipeline de Minutas
Data: Janeiro 2026
Versao: 2.0 - Migrado para google.genai SDK
"""

import argparse
import base64
import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

# Adiciona o diretorio raiz ao path para imports
ROOT_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT_DIR))

# Bibliotecas externas
try:
    from dotenv import load_dotenv
    from google import genai
    from google.genai import types
    from PIL import Image
    import fitz  # PyMuPDF
except ImportError as e:
    print(f"Erro: Biblioteca nao encontrada - {e}")
    print("Instale as dependencias: pip install python-dotenv google-genai Pillow PyMuPDF")
    sys.exit(1)

# Bibliotecas opcionais para conversao DOCX
DOCX_SUPPORT_AVAILABLE = False
DOCX_CONVERSION_METHOD = None

try:
    import docx2pdf
    DOCX_SUPPORT_AVAILABLE = True
    DOCX_CONVERSION_METHOD = 'docx2pdf'
except ImportError:
    pass

# Fallback: tenta usar python-docx + reportlab para conversao manual
if not DOCX_SUPPORT_AVAILABLE:
    try:
        from docx import Document as DocxDocument
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        DOCX_SUPPORT_AVAILABLE = True
        DOCX_CONVERSION_METHOD = 'python-docx'
    except ImportError:
        pass

if not DOCX_SUPPORT_AVAILABLE:
    logging.getLogger(__name__).warning(
        "Suporte a DOCX nao disponivel. Instale: pip install docx2pdf ou pip install python-docx reportlab"
    )

# Carrega variaveis de ambiente
load_dotenv(ROOT_DIR / '.env')

# Configuracao de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Constantes
GEMINI_MODEL = "gemini-3-flash-preview"  # Upgrade para Gemini 3 Flash
PROMPTS_DIR = ROOT_DIR / 'execution' / 'prompts'
TMP_DIR = ROOT_DIR / '.tmp'
RATE_LIMIT_DELAY = 4  # Segundos entre requisicoes (15 RPM = 4s)
MAX_RETRIES = 3
SAVE_PROGRESS_INTERVAL = 5
SUPPORTED_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
SUPPORTED_PDF_EXTENSION = '.pdf'
SUPPORTED_DOCX_EXTENSION = '.docx'


# =============================================================================
# CLASSES DE ERRO
# =============================================================================

class GeminiExtractionError(Exception):
    """Erro durante extracao com Gemini."""
    pass


class PromptNotFoundError(GeminiExtractionError):
    """Prompt nao encontrado para o tipo de documento."""
    pass


class FileLoadError(GeminiExtractionError):
    """Erro ao carregar arquivo."""
    pass


class ParseError(GeminiExtractionError):
    """Erro ao fazer parsing da resposta do Gemini."""
    pass


# =============================================================================
# FUNCOES DE CONFIGURACAO
# =============================================================================

def load_environment() -> str:
    """
    Carrega e valida a API key do Gemini.

    Returns:
        API key do Gemini

    Raises:
        ValueError: Se a API key nao estiver configurada
    """
    api_key = os.getenv('GEMINI_API_KEY')
    if not api_key:
        raise ValueError(
            "GEMINI_API_KEY nao encontrada no ambiente.\n"
            "Configure no arquivo .env"
        )
    return api_key


def configure_gemini(api_key: str) -> genai.Client:
    """
    Configura o cliente Gemini usando o novo SDK google.genai.

    Args:
        api_key: Chave da API

    Returns:
        Cliente Gemini configurado
    """
    # Novo SDK usa Client centralizado
    client = genai.Client(api_key=api_key)
    logger.info(f"Modelo configurado: {GEMINI_MODEL}")
    return client


# =============================================================================
# FUNCOES DE CARREGAMENTO DE PROMPTS
# =============================================================================

def load_prompt(tipo_documento: str) -> str:
    """
    Carrega prompt especifico para o tipo de documento.

    Args:
        tipo_documento: Tipo do documento (ex: MATRICULA_IMOVEL, RG)

    Returns:
        Texto do prompt

    Raises:
        PromptNotFoundError: Se prompt nao existir e generic tambem nao
    """
    # Normaliza o nome do arquivo
    prompt_name = tipo_documento.lower()
    prompt_file = PROMPTS_DIR / f"{prompt_name}.txt"

    if prompt_file.exists():
        content = prompt_file.read_text(encoding='utf-8')
        logger.debug(f"Prompt carregado: {prompt_file.name}")
        return content

    # Tenta carregar prompt generico
    generic_file = PROMPTS_DIR / "generic.txt"
    if generic_file.exists():
        content = generic_file.read_text(encoding='utf-8')
        logger.warning(f"Prompt especifico nao encontrado para {tipo_documento}, usando generic")
        return content

    raise PromptNotFoundError(
        f"Prompt nao encontrado: {prompt_file}\n"
        f"Crie o arquivo ou use o prompt generico."
    )


def list_available_prompts() -> List[str]:
    """
    Lista todos os prompts disponiveis.

    Returns:
        Lista de tipos de documento com prompts
    """
    if not PROMPTS_DIR.exists():
        return []

    prompts = []
    for file in PROMPTS_DIR.glob("*.txt"):
        tipo = file.stem.upper()
        if tipo != "GENERIC":
            prompts.append(tipo)

    return sorted(prompts)


# =============================================================================
# FUNCOES DE CARREGAMENTO DE ARQUIVOS
# =============================================================================

def extract_first_page_from_pdf(pdf_path: Path) -> Optional[Image.Image]:
    """
    Extrai a primeira pagina de um PDF como imagem.
    DEPRECATED: Use extract_all_pages_from_pdf para processar todas as paginas.

    Args:
        pdf_path: Caminho para o arquivo PDF

    Returns:
        Imagem PIL da primeira pagina ou None
    """
    try:
        doc = fitz.open(str(pdf_path))
        if len(doc) == 0:
            logger.warning(f"PDF vazio: {pdf_path}")
            return None

        # Pega a primeira pagina
        page = doc[0]

        # Renderiza em alta resolucao (300 DPI)
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)

        # Converte para PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        doc.close()
        return img

    except Exception as e:
        logger.error(f"Erro ao extrair pagina do PDF {pdf_path}: {e}")
        return None


def extract_all_pages_from_pdf(pdf_path: Path, max_pages: int = 50) -> Optional[Image.Image]:
    """
    Extrai TODAS as paginas de um PDF e concatena em uma unica imagem vertical.

    Esta funcao processa documentos multipagina (como matriculas de imovel)
    criando uma imagem longa que contem todas as paginas empilhadas verticalmente.
    O Gemini consegue processar imagens grandes, entao essa abordagem e preferida.

    Args:
        pdf_path: Caminho para o arquivo PDF
        max_pages: Numero maximo de paginas a processar (padrao: 50)

    Returns:
        Imagem PIL concatenada verticalmente com todas as paginas, ou None se falhar
    """
    try:
        doc = fitz.open(str(pdf_path))
        num_pages = len(doc)

        if num_pages == 0:
            logger.warning(f"PDF vazio: {pdf_path}")
            doc.close()
            return None

        # Limita o numero de paginas se necessario
        pages_to_process = min(num_pages, max_pages)
        if num_pages > max_pages:
            logger.warning(f"PDF tem {num_pages} paginas, processando apenas as primeiras {max_pages}")

        logger.info(f"Extraindo {pages_to_process} pagina(s) do PDF: {pdf_path.name}")

        # Lista para armazenar as imagens de cada pagina
        page_images: List[Image.Image] = []

        # Renderiza cada pagina
        # Usa zoom 2.0 para ~150 DPI (balanco entre qualidade e tamanho)
        # Para PDFs muito grandes, podemos reduzir o zoom
        zoom = 2.0 if pages_to_process <= 10 else 1.5
        mat = fitz.Matrix(zoom, zoom)

        for page_num in range(pages_to_process):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=mat)

            # Converte para PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_images.append(img)

            logger.debug(f"  Pagina {page_num + 1}/{pages_to_process}: {pix.width}x{pix.height} pixels")

        doc.close()

        # Se so tem uma pagina, retorna diretamente
        if len(page_images) == 1:
            return page_images[0]

        # Concatena todas as paginas verticalmente
        # Calcula dimensoes da imagem final
        total_width = max(img.width for img in page_images)
        total_height = sum(img.height for img in page_images)

        # Adiciona uma pequena margem entre as paginas (5 pixels)
        margin = 5
        total_height += margin * (len(page_images) - 1)

        logger.info(f"Concatenando {len(page_images)} paginas em imagem {total_width}x{total_height}")

        # Cria imagem final com fundo branco
        final_image = Image.new('RGB', (total_width, total_height), color=(255, 255, 255))

        # Cola cada pagina na posicao correta
        y_offset = 0
        for idx, img in enumerate(page_images):
            # Centraliza horizontalmente se a pagina for menor que a largura maxima
            x_offset = (total_width - img.width) // 2
            final_image.paste(img, (x_offset, y_offset))
            y_offset += img.height + margin

        # Limpa memoria das imagens individuais
        for img in page_images:
            img.close()

        return final_image

    except Exception as e:
        logger.error(f"Erro ao extrair paginas do PDF {pdf_path}: {e}")
        return None


def convert_docx_to_images(docx_path: Path, max_pages: int = 50) -> Optional[Image.Image]:
    """
    Converte um arquivo DOCX para imagem(s), usando conversao para PDF intermediario.

    Metodos de conversao (em ordem de preferencia):
    1. docx2pdf - Usa Microsoft Word (Windows) ou LibreOffice (Linux/Mac)
    2. python-docx + reportlab - Conversao manual (fallback, qualidade inferior)

    Args:
        docx_path: Caminho para o arquivo DOCX
        max_pages: Numero maximo de paginas a processar

    Returns:
        Imagem PIL concatenada com todas as paginas, ou None se falhar
    """
    import tempfile
    import shutil

    if not DOCX_SUPPORT_AVAILABLE:
        logger.error(f"Suporte a DOCX nao disponivel. Arquivo ignorado: {docx_path}")
        return None

    logger.info(f"Convertendo DOCX para imagem: {docx_path.name} (metodo: {DOCX_CONVERSION_METHOD})")

    try:
        # Cria diretorio temporario para conversao
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            temp_pdf = temp_path / f"{docx_path.stem}.pdf"

            # Metodo 1: docx2pdf (preferido - melhor qualidade)
            if DOCX_CONVERSION_METHOD == 'docx2pdf':
                try:
                    docx2pdf.convert(str(docx_path), str(temp_pdf))

                    if temp_pdf.exists():
                        logger.info(f"DOCX convertido para PDF via docx2pdf")
                        # Usa a funcao existente para converter PDF para imagem
                        return extract_all_pages_from_pdf(temp_pdf, max_pages)
                    else:
                        logger.error(f"docx2pdf nao gerou o arquivo PDF esperado")
                        return None

                except Exception as e:
                    logger.error(f"Erro na conversao docx2pdf: {e}")
                    # Tenta fallback se disponivel
                    if 'DocxDocument' in dir():
                        logger.info("Tentando fallback com python-docx...")
                    else:
                        return None

            # Metodo 2: python-docx + reportlab (fallback)
            if DOCX_CONVERSION_METHOD == 'python-docx':
                try:
                    from docx import Document as DocxDocument
                    from reportlab.lib.pagesizes import A4
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.lib.enums import TA_JUSTIFY

                    # Abre o documento DOCX
                    doc = DocxDocument(str(docx_path))

                    # Configura o PDF
                    pdf_doc = SimpleDocTemplate(
                        str(temp_pdf),
                        pagesize=A4,
                        rightMargin=72,
                        leftMargin=72,
                        topMargin=72,
                        bottomMargin=72
                    )

                    # Estilos
                    styles = getSampleStyleSheet()
                    normal_style = ParagraphStyle(
                        'CustomNormal',
                        parent=styles['Normal'],
                        fontSize=10,
                        leading=14,
                        alignment=TA_JUSTIFY
                    )

                    # Constroi o conteudo
                    story = []

                    for para in doc.paragraphs:
                        if para.text.strip():
                            # Escapa caracteres especiais do XML
                            text = para.text.replace('&', '&amp;')
                            text = text.replace('<', '&lt;')
                            text = text.replace('>', '&gt;')
                            story.append(Paragraph(text, normal_style))
                            story.append(Spacer(1, 6))

                    # Processa tabelas
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                row_text = row_text.replace('&', '&amp;')
                                row_text = row_text.replace('<', '&lt;')
                                row_text = row_text.replace('>', '&gt;')
                                story.append(Paragraph(row_text, normal_style))
                                story.append(Spacer(1, 3))

                    if not story:
                        logger.warning(f"DOCX vazio ou sem conteudo extraivel: {docx_path}")
                        return None

                    # Gera o PDF
                    pdf_doc.build(story)

                    if temp_pdf.exists():
                        logger.info(f"DOCX convertido para PDF via python-docx (fallback)")
                        return extract_all_pages_from_pdf(temp_pdf, max_pages)
                    else:
                        logger.error(f"Falha ao gerar PDF com python-docx")
                        return None

                except Exception as e:
                    logger.error(f"Erro na conversao python-docx: {e}")
                    return None

    except Exception as e:
        logger.error(f"Erro ao converter DOCX {docx_path}: {e}")
        return None

    return None


def load_original_file(file_path: Path) -> Tuple[Optional[bytes], Optional[str]]:
    """
    Carrega arquivo original como bytes e detecta MIME type.

    Args:
        file_path: Caminho para o arquivo

    Returns:
        Tuple (bytes do arquivo, MIME type) ou (None, None) se falhar
    """
    if not file_path.exists():
        logger.error(f"Arquivo nao encontrado: {file_path}")
        return None, None

    ext = file_path.suffix.lower()

    # Mapeamento de extensao para MIME type
    # Nota: .docx requer conversao especial, nao tem MIME direto para envio ao Gemini
    mime_types = {
        '.pdf': 'application/pdf',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    mime_type = mime_types.get(ext)
    if not mime_type:
        logger.warning(f"Extensao nao suportada: {ext}")
        return None, None

    try:
        # Para arquivos DOCX, converte para imagem via PDF intermediario
        if ext == '.docx':
            if not DOCX_SUPPORT_AVAILABLE:
                logger.error(f"Suporte a DOCX nao disponivel. Instale: pip install docx2pdf")
                return None, None

            img = convert_docx_to_images(file_path)
            if img is None:
                logger.error(f"Falha ao converter DOCX: {file_path}")
                return None, None

            # Converte imagem para bytes JPEG (mesmo processo do PDF)
            import io
            buffer = io.BytesIO()

            # Calcula qualidade baseada no tamanho
            total_pixels = img.width * img.height
            if total_pixels > 50_000_000:
                quality = 70
            elif total_pixels > 20_000_000:
                quality = 80
            else:
                quality = 95

            img.save(buffer, format='JPEG', quality=quality)
            image_bytes = buffer.getvalue()

            size_mb = len(image_bytes) / (1024 * 1024)
            logger.info(f"DOCX convertido: {img.width}x{img.height} pixels, {size_mb:.2f}MB")

            img.close()
            return image_bytes, 'image/jpeg'

        # Para PDFs, converte TODAS as paginas para imagem concatenada
        if ext == '.pdf':
            img = extract_all_pages_from_pdf(file_path)
            if img is None:
                return None, None

            # Converte imagem para bytes JPEG
            # Usa qualidade adaptativa baseada no tamanho da imagem
            import io
            buffer = io.BytesIO()

            # Calcula qualidade baseada no tamanho total (para evitar imagens muito grandes)
            # Gemini aceita ate ~20MB, mas queremos manter razoavel
            total_pixels = img.width * img.height
            if total_pixels > 50_000_000:  # Mais de 50 megapixels
                quality = 70
                logger.info(f"Imagem muito grande ({total_pixels/1_000_000:.1f}MP), usando qualidade {quality}%")
            elif total_pixels > 20_000_000:  # Mais de 20 megapixels
                quality = 80
            else:
                quality = 95

            img.save(buffer, format='JPEG', quality=quality)
            image_bytes = buffer.getvalue()

            # Log do tamanho final
            size_mb = len(image_bytes) / (1024 * 1024)
            logger.info(f"PDF convertido: {img.width}x{img.height} pixels, {size_mb:.2f}MB")

            img.close()
            return image_bytes, 'image/jpeg'

        # Para imagens, le diretamente
        with open(file_path, 'rb') as f:
            content = f.read()

        return content, mime_type

    except Exception as e:
        logger.error(f"Erro ao carregar arquivo {file_path}: {e}")
        return None, None


def load_ocr_text(ocr_path: str) -> Optional[str]:
    """
    Carrega texto OCR correspondente.

    Args:
        ocr_path: Caminho para o arquivo OCR

    Returns:
        Texto OCR ou None
    """
    try:
        path = Path(ocr_path)
        if not path.exists():
            logger.warning(f"Arquivo OCR nao encontrado: {ocr_path}")
            return None

        content = path.read_text(encoding='utf-8')

        # Remove cabecalho do OCR (linhas antes de ---)
        if '---' in content:
            content = content.split('---', 1)[1]

        return content.strip()

    except Exception as e:
        logger.error(f"Erro ao carregar OCR {ocr_path}: {e}")
        return None


# =============================================================================
# FUNCOES DE CHAMADA AO GEMINI
# =============================================================================

def call_gemini(
    client: genai.Client,
    prompt: str,
    image_bytes: bytes,
    mime_type: str,
    ocr_text: Optional[str] = None
) -> str:
    """
    Chama Gemini com imagem usando o novo SDK google.genai.

    Args:
        client: Cliente Gemini configurado
        prompt: Prompt do tipo de documento
        image_bytes: Bytes da imagem
        mime_type: Tipo MIME da imagem
        ocr_text: Texto OCR (opcional, para referencia - DEPRECATED)

    Returns:
        Texto da resposta do Gemini
    """
    # Monta o prompt completo (OCR nao e mais usado, mas mantido para compatibilidade)
    full_prompt = prompt

    # Prepara a imagem no formato do novo SDK
    image_part = types.Part.from_bytes(
        data=image_bytes,
        mime_type=mime_type
    )

    # Configura geracao com temperatura baixa para extracao precisa
    config = types.GenerateContentConfig(
        temperature=0.1,
        max_output_tokens=16384
    )

    # Envia para o Gemini usando o novo SDK
    # Tenta usar gemini-3-flash-preview com fallbacks
    model_names = [GEMINI_MODEL, "gemini-2.5-flash", "gemini-2.0-flash"]

    last_error = None
    for model_name in model_names:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=[image_part, full_prompt],
                config=config
            )

            # Validacao: verifica se response tem texto
            if response is None or not response.text:
                finish_reason = getattr(response, 'finish_reason', 'UNKNOWN') if response else 'NO_RESPONSE'
                logger.warning(
                    f"Modelo {model_name} retornou resposta vazia. Finish reason: {finish_reason}"
                )
                # Tenta proximo modelo
                continue

            return response.text
        except Exception as e:
            last_error = e
            if "not found" in str(e).lower() or "unavailable" in str(e).lower():
                logger.warning(f"Modelo {model_name} nao disponivel, tentando proximo...")
                continue
            raise  # Re-raise se for outro tipo de erro

    raise RuntimeError(f"Nenhum modelo Gemini disponivel: {last_error}")


def call_gemini_with_retry(
    client: genai.Client,
    prompt: str,
    image_bytes: bytes,
    mime_type: str,
    ocr_text: Optional[str] = None,
    max_retries: int = MAX_RETRIES
) -> str:
    """
    Chama Gemini com retry e backoff exponencial.

    Args:
        client: Cliente Gemini configurado
        prompt: Prompt do tipo de documento
        image_bytes: Bytes da imagem
        mime_type: Tipo MIME da imagem
        ocr_text: Texto OCR (opcional - DEPRECATED)
        max_retries: Numero maximo de tentativas

    Returns:
        Texto da resposta do Gemini

    Raises:
        GeminiExtractionError: Se todas as tentativas falharem
    """
    last_error = None

    for attempt in range(max_retries):
        try:
            return call_gemini(client, prompt, image_bytes, mime_type, ocr_text)

        except Exception as e:
            last_error = e
            wait_time = (2 ** attempt) * 2  # 2, 4, 8 segundos

            logger.warning(f"Tentativa {attempt + 1}/{max_retries} falhou: {e}")

            if attempt < max_retries - 1:
                logger.info(f"Aguardando {wait_time}s antes de retry...")
                time.sleep(wait_time)

    raise GeminiExtractionError(f"Falha apos {max_retries} tentativas: {last_error}")


# =============================================================================
# FUNCOES DE PARSING DA RESPOSTA
# =============================================================================

def parse_gemini_response(response: str) -> Dict[str, Any]:
    """
    Parseia resposta do Gemini em 3 partes.

    Args:
        response: Texto da resposta do Gemini

    Returns:
        Dicionario com reescrita, explicacao e dados_catalogados
    """
    result = {
        "reescrita": "",
        "explicacao": "",
        "dados_catalogados": {}
    }

    # Valida se response e None ou vazio
    if not response:
        logger.warning("Resposta do Gemini vazia ou None")
        return result

    # Regex patterns para extrair secoes
    patterns = {
        "reescrita": r"##\s*REESCRITA\s+DO\s+DOCUMENTO\s*\n(.*?)(?=##\s*EXPLICA[CÇ][AÃ]O|$)",
        "explicacao": r"##\s*EXPLICA[CÇ][AÃ]O\s+CONTEXTUAL\s*\n(.*?)(?=##\s*DADOS|$)",
        "json": r"```json\s*\n(.*?)\n```"
    }

    # Extrai reescrita
    match = re.search(patterns["reescrita"], response, re.DOTALL | re.IGNORECASE)
    if match:
        result["reescrita"] = match.group(1).strip()

    # Extrai explicacao
    match = re.search(patterns["explicacao"], response, re.DOTALL | re.IGNORECASE)
    if match:
        result["explicacao"] = match.group(1).strip()

    # Extrai JSON
    match = re.search(patterns["json"], response, re.DOTALL)
    if match:
        try:
            json_str = match.group(1).strip()
            result["dados_catalogados"] = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.warning(f"Erro ao fazer parsing do JSON: {e}")
            # Tenta limpar e parsear novamente
            try:
                # Remove caracteres de controle
                json_str = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
                result["dados_catalogados"] = json.loads(json_str)
            except json.JSONDecodeError:
                result["dados_catalogados"] = {"_raw_json": json_str, "_parse_error": str(e)}

    return result


# =============================================================================
# FUNCOES DE PROCESSAMENTO
# =============================================================================

def process_document(
    client: genai.Client,
    doc_info: Dict[str, Any],
    escritura_id: str
) -> Dict[str, Any]:
    """
    Processa um documento completo com Gemini.

    Args:
        client: Cliente Gemini configurado
        doc_info: Informacoes do documento do catalogo
        escritura_id: ID da escritura

    Returns:
        Dicionario com resultado da extracao
    """
    start_time = time.time()

    result = {
        "tipo_documento": doc_info.get('tipo_documento', 'OUTRO'),
        "arquivo_origem": doc_info.get('nome', ''),
        "arquivo_ocr": doc_info.get('arquivo_ocr', ''),
        "data_processamento": datetime.now().isoformat(),
        "modelo": GEMINI_MODEL,
        "reescrita_interpretada": "",
        "explicacao_contextual": "",
        "dados_catalogados": {},
        "metadados": {
            "tokens_entrada": 0,
            "tokens_saida": 0,
            "tempo_processamento_s": 0
        },
        "status": "sucesso",
        "erro": None
    }

    try:
        # 1. Carrega o arquivo original
        file_path = Path(doc_info.get('caminho_absoluto', ''))
        image_bytes, mime_type = load_original_file(file_path)

        if image_bytes is None:
            raise FileLoadError(f"Nao foi possivel carregar: {file_path}")

        # 2. Carrega texto OCR (opcional)
        ocr_text = None
        if doc_info.get('arquivo_ocr'):
            ocr_text = load_ocr_text(doc_info['arquivo_ocr'])

        # 3. Carrega prompt especifico
        tipo = doc_info.get('tipo_documento', 'OUTRO')
        prompt = load_prompt(tipo)

        # 4. Chama Gemini
        response = call_gemini_with_retry(
            client=client,
            prompt=prompt,
            image_bytes=image_bytes,
            mime_type=mime_type,
            ocr_text=ocr_text
        )

        # 5. Parseia resposta
        parsed = parse_gemini_response(response)

        result["reescrita_interpretada"] = parsed["reescrita"]
        result["explicacao_contextual"] = parsed["explicacao"]
        result["dados_catalogados"] = parsed["dados_catalogados"]

        # Estima tokens (aproximacao)
        result["metadados"]["tokens_entrada"] = len(prompt.split()) + (len(ocr_text.split()) if ocr_text else 0)
        result["metadados"]["tokens_saida"] = len(response.split()) if response else 0

    except PromptNotFoundError as e:
        result["status"] = "erro"
        result["erro"] = f"Prompt nao encontrado: {e}"
        logger.error(f"Prompt nao encontrado para {doc_info.get('nome')}: {e}")

    except FileLoadError as e:
        result["status"] = "erro"
        result["erro"] = f"Erro ao carregar arquivo: {e}"
        logger.error(f"Erro ao carregar {doc_info.get('nome')}: {e}")

    except GeminiExtractionError as e:
        result["status"] = "erro"
        result["erro"] = f"Erro na extracao Gemini: {e}"
        logger.error(f"Erro Gemini para {doc_info.get('nome')}: {e}")

    except Exception as e:
        result["status"] = "erro"
        result["erro"] = f"Erro inesperado: {e}"
        logger.exception(f"Erro inesperado ao processar {doc_info.get('nome')}")

    finally:
        result["metadados"]["tempo_processamento_s"] = round(time.time() - start_time, 2)

    return result


def load_catalog(escritura_id: str) -> Dict[str, Any]:
    """
    Carrega o catalogo de uma escritura.

    Args:
        escritura_id: ID da escritura

    Returns:
        Dados do catalogo

    Raises:
        FileNotFoundError: Se catalogo nao existir
    """
    catalog_path = TMP_DIR / 'catalogos' / f'{escritura_id}.json'

    if not catalog_path.exists():
        raise FileNotFoundError(f"Catalogo nao encontrado: {catalog_path}")

    with open(catalog_path, 'r', encoding='utf-8') as f:
        catalog = json.load(f)

    logger.info(f"Catalogo carregado: {len(catalog.get('arquivos', []))} arquivos")
    return catalog


def run_extraction(
    escritura_id: str,
    limit: Optional[int] = None,
    tipo_filtro: Optional[str] = None,
    verbose: bool = False
) -> Dict[str, Any]:
    """
    Executa extracao contextual para toda a escritura.

    Args:
        escritura_id: ID da escritura
        limit: Limite de arquivos a processar
        tipo_filtro: Filtrar por tipo de documento
        verbose: Modo verbose

    Returns:
        Estatisticas e resultados da extracao
    """
    start_time = datetime.now()

    # Configura Gemini
    api_key = load_environment()
    client = configure_gemini(api_key)

    # Carrega catalogo
    catalog = load_catalog(escritura_id)

    # Filtra arquivos
    arquivos = catalog.get('arquivos', [])

    # Filtra apenas arquivos com OCR processado
    arquivos = [a for a in arquivos if a.get('status_ocr') == 'processado']

    # Filtra por tipo se especificado
    if tipo_filtro:
        tipo_upper = tipo_filtro.upper()
        arquivos = [a for a in arquivos if a.get('tipo_documento') == tipo_upper]

    # Aplica limite
    if limit:
        arquivos = arquivos[:limit]

    total = len(arquivos)

    if total == 0:
        logger.warning("Nenhum arquivo para processar")
        return {
            'escritura_id': escritura_id,
            'data_extracao': start_time.isoformat(),
            'total_arquivos': 0,
            'extraidos_sucesso': 0,
            'extraidos_erro': 0,
            'extracoes': []
        }

    logger.info("=" * 60)
    logger.info("EXTRACAO CONTEXTUAL COM GEMINI")
    logger.info(f"  Escritura: {escritura_id}")
    logger.info(f"  Total de arquivos: {total}")
    logger.info(f"  Modelo: {GEMINI_MODEL}")
    if tipo_filtro:
        logger.info(f"  Filtro de tipo: {tipo_filtro}")
    logger.info(f"  Prompts disponiveis: {len(list_available_prompts())}")
    logger.info("=" * 60)

    # Diretorio de saida
    output_dir = TMP_DIR / 'contextual' / escritura_id
    output_dir.mkdir(parents=True, exist_ok=True)

    # Resultados
    extracoes = []
    sucesso = 0
    erro = 0

    for idx, arquivo in enumerate(arquivos, 1):
        nome = arquivo['nome']
        tipo = arquivo.get('tipo_documento', 'OUTRO')

        logger.info(f"[{idx}/{total}] Processando: {nome} ({tipo})")

        # Processa documento
        resultado = process_document(client, arquivo, escritura_id)

        # Adiciona metadados do catalogo
        resultado['id'] = arquivo['id']
        resultado['nome_arquivo'] = nome
        resultado['pessoa_relacionada'] = arquivo.get('pessoa_relacionada')
        resultado['papel_inferido'] = arquivo.get('papel_inferido')

        # Atualiza contadores
        if resultado['status'] == 'sucesso':
            sucesso += 1
            if verbose:
                logger.info(f"    Campos extraidos: {len(resultado.get('dados_catalogados', {}))}")
                logger.info(f"    Tempo: {resultado['metadados']['tempo_processamento_s']}s")
        else:
            erro += 1
            logger.warning(f"    Erro: {resultado.get('erro', 'desconhecido')}")

        extracoes.append(resultado)

        # Salva resultado individual
        output_file = output_dir / f"{arquivo['id']}_{tipo}.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(resultado, f, ensure_ascii=False, indent=2)

        # Salva progresso
        if idx % SAVE_PROGRESS_INTERVAL == 0:
            logger.info(f"  Progresso: {idx}/{total} arquivos processados")

        # Rate limiting
        if idx < total:
            logger.debug(f"Rate limiting: aguardando {RATE_LIMIT_DELAY}s...")
            time.sleep(RATE_LIMIT_DELAY)

    # Calcula estatisticas
    tempos = [e['metadados']['tempo_processamento_s'] for e in extracoes]
    tempo_medio = sum(tempos) / len(tempos) if tempos else 0

    # Resultado final
    resultado_final = {
        'escritura_id': escritura_id,
        'data_extracao': start_time.isoformat(),
        'data_conclusao': datetime.now().isoformat(),
        'tempo_processamento_total': (datetime.now() - start_time).total_seconds(),
        'tempo_medio_por_documento': round(tempo_medio, 2),
        'modelo': GEMINI_MODEL,
        'total_arquivos': total,
        'extraidos_sucesso': sucesso,
        'extraidos_erro': erro,
        'taxa_sucesso': round(sucesso / total * 100, 1) if total > 0 else 0,
        'extracoes': extracoes
    }

    # Salva relatorio completo
    report_path = output_dir / 'relatorio_contextual.json'
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(resultado_final, f, ensure_ascii=False, indent=2)

    # Resumo final
    logger.info("=" * 60)
    logger.info("EXTRACAO CONTEXTUAL CONCLUIDA")
    logger.info(f"  Escritura: {escritura_id}")
    logger.info(f"  Total processados: {total}")
    logger.info(f"  Sucesso: {sucesso}")
    logger.info(f"  Erros: {erro}")
    logger.info(f"  Taxa de sucesso: {resultado_final['taxa_sucesso']}%")
    logger.info(f"  Tempo total: {resultado_final['tempo_processamento_total']:.2f}s")
    logger.info(f"  Tempo medio/doc: {tempo_medio:.2f}s")
    logger.info(f"  Saida: {output_dir}")
    logger.info("=" * 60)

    return resultado_final


def main():
    """Funcao principal - entry point do script"""
    parser = argparse.ArgumentParser(
        description='Extracao contextual de documentos com Gemini 3 Flash (SDK google.genai)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python extract_with_gemini.py FC_515_124_p280509
  python extract_with_gemini.py FC_515_124_p280509 --type MATRICULA_IMOVEL
  python extract_with_gemini.py FC_515_124_p280509 --limit 5
  python extract_with_gemini.py FC_515_124_p280509 --verbose

O script usa a visao multimodal do Gemini 3 Flash para interpretar documentos
de cartorio DIRETAMENTE (sem OCR), gerando reescrita organizada, explicacao
contextual e dados estruturados.

Prompts disponiveis em: execution/prompts/
        """
    )

    parser.add_argument(
        'escritura_id',
        nargs='?',
        default=None,
        help='ID da escritura (ex: FC_515_124_p280509)'
    )

    parser.add_argument(
        '--limit', '-l',
        type=int,
        help='Limita o numero de arquivos processados'
    )

    parser.add_argument(
        '--type', '-t',
        type=str,
        dest='tipo',
        help='Filtra por tipo de documento (ex: MATRICULA_IMOVEL, RG, CNDT)'
    )

    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Modo verbose: mostra mais detalhes'
    )

    parser.add_argument(
        '--list-prompts',
        action='store_true',
        help='Lista prompts disponiveis e sai'
    )

    args = parser.parse_args()

    # Configura nivel de log
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Lista prompts se solicitado
    if args.list_prompts:
        print("\nPrompts disponiveis:")
        print("-" * 40)
        for prompt in list_available_prompts():
            print(f"  - {prompt}")
        print("-" * 40)
        print(f"Total: {len(list_available_prompts())}")
        print(f"Diretorio: {PROMPTS_DIR}")
        sys.exit(0)

    # Valida que escritura_id foi fornecido
    if not args.escritura_id:
        parser.error("escritura_id e obrigatorio (ex: FC_515_124_p280509)")

    try:
        result = run_extraction(
            escritura_id=args.escritura_id,
            limit=args.limit,
            tipo_filtro=args.tipo,
            verbose=args.verbose
        )

        # Retorna codigo de saida baseado em erros
        if result['extraidos_erro'] > result['extraidos_sucesso']:
            sys.exit(2)  # Mais erros que sucesso
        sys.exit(0)

    except FileNotFoundError as e:
        logger.error(f"Arquivo nao encontrado: {e}")
        sys.exit(1)
    except ValueError as e:
        logger.error(f"Erro de configuracao: {e}")
        sys.exit(1)
    except Exception as e:
        logger.exception(f"Erro fatal: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
